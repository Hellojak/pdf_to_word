import os
import sys
import argparse
import tempfile
import logging
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict
from configparser import ConfigParser

import fitz  # PyMuPDF
from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt
from PIL import Image

__version__ = "2.1.0"

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('pdf_to_word.log'),
        logging.StreamHandler()
    ]
)

class PDFConversionError(Exception):
    pass

class PDFConverter:
    def __init__(self, config_path: str = "config.ini"):
        self.config = self._load_config(config_path)
        self._validate_dependencies()
        
    def _load_config(self, path: str) -> ConfigParser:
        """Load configuration from INI file"""
        config = ConfigParser()
        if Path(path).exists():
            config.read(path)
        else:
            config['PATHS'] = {
                'poppler': '',
                'tesseract': ''
            }
            config['SETTINGS'] = {
                'default_dpi': '300',
                'ocr_languages': 'eng',
                'log_level': 'INFO'
            }
        return config

    def _validate_dependencies(self) -> None:
        """Validate required external dependencies"""
        # Validate Tesseract
        tesseract_path = self.config.get('PATHS', 'tesseract', fallback='')
        if tesseract_path and Path(tesseract_path).exists():
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        elif not pytesseract.get_tesseract_version():
            raise EnvironmentError("Tesseract OCR not found. Install or specify path in config.ini")

        # Validate Poppler
        self.poppler_path = self.config.get('PATHS', 'poppler', fallback=None)
        if self.poppler_path and not Path(self.poppler_path).exists():
            raise FileNotFoundError(f"Poppler not found at: {self.poppler_path}")

    def convert(
        self,
        input_pdf: str,
        output_docx: str,
        dpi: int = None,
        password: str = None,
        verbose: bool = False
    ) -> None:
        """Main conversion method with enhanced features"""
        self._validate_input(input_pdf, output_docx)
        dpi = dpi or int(self.config.get('SETTINGS', 'default_dpi', fallback=300))
        
        try:
            if verbose:
                logging.getLogger().setLevel(logging.DEBUG)

            logging.info(f"Starting conversion (v{__version__}): {input_pdf}")
            
            if self._is_scanned_pdf(input_pdf, password):
                self._convert_with_ocr(input_pdf, output_docx, dpi)
            else:
                self._convert_with_styles(input_pdf, output_docx, password)
                
            logging.info(f"Successfully created: {output_docx}")

        except Exception as e:
            logging.error(f"Conversion failed: {str(e)}")
            raise PDFConversionError(f"Failed to convert {input_pdf}") from e

    def _validate_input(self, input_pdf: str, output_docx: str) -> None:
        """Validate input parameters"""
        if not Path(input_pdf).exists():
            raise FileNotFoundError(f"Input PDF not found: {input_pdf}")
            
        if Path(output_docx).exists():
            raise FileExistsError(f"Output file already exists: {output_docx}")

    def _is_scanned_pdf(self, file_path: str, password: str = None) -> bool:
        """Determine if PDF is image-based"""
        try:
            with fitz.open(file_path) as doc:
                if doc.is_encrypted and password:
                    doc.authenticate(password)
                    
                for page in doc:
                    if page.get_text().strip():
                        return False
                return True
        except Exception as e:
            logging.error(f"PDF analysis failed: {str(e)}")
            return True

    def _convert_with_styles(
        self,
        input_pdf: str,
        output_docx: str,
        password: str = None
    ) -> None:
        """Convert text-based PDF with style preservation"""
        doc = Document()
        try:
            with fitz.open(input_pdf) as pdf:
                if pdf.is_encrypted:
                    if password:
                        pdf.authenticate(password)
                    else:
                        raise PDFConversionError("PDF is encrypted - password required")

                logging.info(f"Processing {len(pdf)} pages with text extraction")
                
                for page in tqdm(pdf, desc="Extracting text pages"):
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if block['type'] == 0:
                            self._process_text_block(doc, block)
                    doc.add_page_break()
            
            doc.save(output_docx)
            logging.debug(f"Saved styled document: {output_docx}")

        except Exception as e:
            logging.error(f"Text extraction failed: {str(e)}")
            raise

    def _process_text_block(self, doc: Document, block: Dict) -> None:
        """Process text block with styling"""
        para = doc.add_paragraph()
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                text = span["text"]
                if text.strip():
                    self._add_styled_run(para, text, span)

    def _add_styled_run(self, para, text: str, span: Dict) -> None:
        """Add text with original formatting"""
        try:
            run = para.add_run(text)
            run.bold = bool(span["flags"] & 16) or "bold" in span["font"].lower()
            run.italic = bool(span["flags"] & 2) or "italic" in span["font"].lower()
            
            # Set font size
            font_size = span["size"]
            run.font.size = Pt(font_size)
            
            logging.debug(f"Added text: {text[:50]}... (Size: {font_size}pt, Bold: {run.bold}, Italic: {run.italic})")
        except Exception as e:
            logging.error(f"Failed to add styled run: {str(e)}")
            raise

    def _convert_with_ocr(self, input_pdf: str, output_docx: str, dpi: int = 300) -> None:
        """Convert scanned PDF using OCR with in-memory processing"""
        doc = Document()
        langs = self.config.get('SETTINGS', 'ocr_languages', fallback='eng')
    
        try:
            logging.info(f"Converting PDF to images (DPI: {dpi})")
            images = convert_from_path(
                input_pdf,
                dpi=dpi,
                poppler_path=self.poppler_path,
                fmt='png',
                thread_count=4,
                use_pdftocairo=True,
                strict=True
            )

            logging.info(f"OCR processing {len(images)} pages")
            for i, image in enumerate(tqdm(images, desc="OCR Processing"), 1):
                with image:  # Ensure proper resource cleanup
                    text = pytesseract.image_to_string(
                        image,
                        lang=langs,
                        config='--psm 6 -c preserve_interword_spaces=1'
                    )
                
                if text.strip():
                    doc.add_paragraph(text)
            
                if i < len(images):
                    doc.add_page_break()

            doc.save(output_docx)
            logging.debug(f"Saved OCR document: {output_docx}")

        except Exception as e:
            logging.error(f"OCR processing failed: {str(e)}")
            raise

def main():
    parser = argparse.ArgumentParser(
        description=f"PDF to Word Converter v{__version__}",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument("input", help="Path to input PDF file")
    parser.add_argument("-o", "--output", help="Custom output DOCX path")
    parser.add_argument("-d", "--dpi", type=int, default=300,
                      help="Image resolution for OCR processing")
    parser.add_argument("-p", "--password", help="Password for encrypted PDF")
    parser.add_argument("-v", "--verbose", action="store_true",
                      help="Enable debug logging")
    
    args = parser.parse_args()

    try:
        # Generate output filename with timestamp
        input_path = Path(args.input)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_output = input_path.with_name(
            f"{input_path.stem}_converted_{timestamp}.docx"
        )
        output_path = Path(args.output) if args.output else default_output

        converter = PDFConverter()
        converter.convert(
            input_pdf=str(input_path.resolve()),
            output_docx=str(output_path.resolve()),
            dpi=args.dpi,
            password=args.password,
            verbose=args.verbose
        )

        print(f"\n✅ Successfully converted to: {output_path}")
        print(f"Check 'pdf_to_word.log' for detailed conversion report")

    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()